"""
app.py — 國立臺北教育大學教務處法規知識庫助理
RAG 架構：LangChain + FAISS + GPT-4o-mini
"""

import os
import re
import tempfile
from pathlib import Path
from dotenv import load_dotenv

import streamlit as st
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import FAISS
from langchain.chains import create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain.memory import ConversationBufferWindowMemory
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
import pdfplumber
import docx

load_dotenv()

# Render 掛載 Disk 路徑（本地開發退回專案目錄）
_DATA_DIR       = Path(os.environ.get("DATA_DIR", "."))
INDEX_DIR       = _DATA_DIR / "faiss_index"
REGULATIONS_DIR = _DATA_DIR / "regulations"
TOP_K          = 5
WINDOW_K       = 5
CHUNK_SIZE     = 500
CHUNK_OVERLAP  = 100

QUICK_QUESTIONS = {
    "student": {
        "label": "在校學生", "icon": "🎓",
        "questions": [
            "如何申請停修課程？", "輔系與雙主修怎麼辦理？",
            "延畢需要什麼條件？", "學分抵免怎麼申請？", "校際選課如何辦理？",
        ],
    },
    "staff": {
        "label": "教職人員／行政", "icon": "🏛️",
        "questions": [
            "教師評鑑的評鑑準則為何？", "教師升等審查的程序？",
            "教學優良獎的評選辦法？", "論文原創性比對聲明書何時實施？",
        ],
    },
    "applicant": {
        "label": "考生／準新生", "icon": "📝",
        "questions": [
            "心理與諮商學系碩士班考哪些科目？", "數學暨資訊教育學系碩士班考科？",
            "新生如何申請保留入學資格？", "轉學生可以申請哪些系？",
        ],
    },
    "graduate": {
        "label": "研究生", "icon": "📚",
        "questions": [
            "學位考試如何申請？", "論文指導費與口試費怎麼支給？",
            "學術倫理時數何時需要繳交？", "論文可以申請延後公開嗎？",
        ],
    },
}

SYSTEM_PROMPT = """你是「國立臺北教育大學（NTUE）教務處數位法規助理」。

【核心原則】
1. 只根據下方「參考法規段落」的內容回答，禁止自行推測或補充段落中未出現的資訊。
2. 若參考段落中有 [QA] 標記的內容，優先以該內容作為主要回答，其他段落作為補充。
3. 回答中不需標註任何來源或出處。
4. 若段落無法完整回答，說明「本系統資料未涵蓋此項目」，建議致電 (02)2732-1104。
5. 使用繁體中文，語氣親切專業。
6. 條列式回答，字數 500 字內。

【參考法規段落】
{context}

【回答】"""


# ── 文件解析工具 ──────────────────────────────────

def clean_text(text: str) -> str:
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"^\s*\d+\s*$", "", text, flags=re.MULTILINE)
    return text.strip()


def is_scanned_pdf(pdf_path: Path) -> bool:
    """偵測是否為掃描版 PDF（文字少於 50 字視為掃描）"""
    with pdfplumber.open(pdf_path) as pdf:
        total = sum(len(p.extract_text() or "") for p in pdf.pages[:3])
    return total < 50


def parse_pdf(file_path: Path, doc_type: str = "regulation") -> list[Document]:
    docs = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = clean_text(page.extract_text() or "")
            if not text:
                continue
            prefix = "[QA] " if doc_type == "qa" else ""
            docs.append(Document(
                page_content=prefix + text,
                metadata={"source": file_path.name, "page": i + 1, "type": doc_type}
            ))
    return docs


def parse_docx(file_path: Path, doc_type: str = "qa") -> list[Document]:
    """解析 DOCX
    QA 格式支援兩種：
      1. 表格格式（大分類 / 小分類 / Q問 / A答）→ 每列一塊
      2. 文字格式（Q: ... A: ...）→ 每組一塊
    """
    d = docx.Document(str(file_path))

    if doc_type != "qa":
        full_text = clean_text("\n".join(p.text for p in d.paragraphs if p.text.strip()))
        return [Document(
            page_content=full_text,
            metadata={"source": file_path.name, "page": 1, "type": doc_type}
        )]

    docs = []

    # 方式 1：表格格式（優先）
    for table in d.tables:
        cols = len(table.columns)
        for row_idx, row in enumerate(table.rows):
            cells = [c.text.strip() for c in row.cells]
            if row_idx == 0:
                continue  # 跳過標題列

            if cols >= 4:
                # 4欄：大分類、小分類、Q問、A答
                big_cat = cells[0] if cells[0] else ""
                sub_cat = cells[1] if len(cells) > 1 else ""
                q_text  = cells[2] if len(cells) > 2 else ""
                a_text  = cells[3] if len(cells) > 3 else ""
            elif cols == 2:
                # 2欄：Q問、A答
                big_cat, sub_cat = "", ""
                q_text = cells[0]
                a_text = cells[1] if len(cells) > 1 else ""
            else:
                continue

            if not q_text or not a_text:
                continue
            # 過濾表格標題列（內容為欄位名稱的列）
            if q_text in ("Q問", "Q", "問題") and a_text in ("A答", "A", "答案"):
                continue

            # 每個 QA 列組成一塊，加上分類 context
            content_parts = []
            if big_cat:
                content_parts.append(f"類別：{big_cat}" + (f" > {sub_cat}" if sub_cat else ""))
            content_parts.append(f"問：{q_text}")
            content_parts.append(f"答：{a_text}")
            block = "\n".join(content_parts)

            docs.append(Document(
                page_content="[QA] " + clean_text(block),
                metadata={
                    "source":   file_path.name,
                    "page":     row_idx,
                    "type":     "qa",
                    "category": big_cat,
                }
            ))

    # 方式 2：純文字 Q: A: 格式（表格為空時使用）
    if not docs:
        full_text = clean_text("\n".join(p.text for p in d.paragraphs if p.text.strip()))
        blocks = re.split(r"(?=^Q[:：])", full_text, flags=re.MULTILINE)
        for block in blocks:
            block = block.strip()
            if not block or len(block) < 10:
                continue
            docs.append(Document(
                page_content="[QA] " + block,
                metadata={"source": file_path.name, "page": 1, "type": "qa"}
            ))

    # 兩種格式都沒有→整份當一塊
    if not docs:
        full_text = clean_text("\n".join(p.text for p in d.paragraphs if p.text.strip()))
        docs = [Document(
            page_content="[QA] " + full_text,
            metadata={"source": file_path.name, "page": 1, "type": "qa"}
        )]

    return docs


def parse_txt(file_path: Path, doc_type: str = "regulation") -> list[Document]:
    text = clean_text(file_path.read_text(encoding="utf-8"))
    prefix = "[QA] " if doc_type == "qa" else ""
    return [Document(
        page_content=prefix + text,
        metadata={"source": file_path.name, "page": 1, "type": doc_type}
    )]


def add_to_index(new_docs: list[Document]) -> tuple[bool, str]:
    """增量更新：只向量化新文件，加進現有 index，不重建全部"""
    if not new_docs:
        return False, "沒有新文件"

    if not INDEX_DIR.exists():
        return False, "找不到現有向量庫，請先執行 build_index.py"

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""],
    )

    # QA 不切塊，法規文件切塊
    qa_docs  = [d for d in new_docs if d.metadata.get("type") == "qa"]
    reg_docs = [d for d in new_docs if d.metadata.get("type") != "qa"]
    chunks   = splitter.split_documents(reg_docs) + qa_docs

    if not chunks:
        return False, "文件解析後無內容"

    embeddings  = OpenAIEmbeddings(model="text-embedding-3-small")

    # 載入現有 index 並增量加入
    vectorstore = FAISS.load_local(
        str(INDEX_DIR), embeddings, allow_dangerous_deserialization=True
    )

    # 分批送，每批 50 塊，避免超過 API token 限制
    BATCH = 50
    for i in range(0, len(chunks), BATCH):
        vectorstore.add_documents(chunks[i:i + BATCH])

    vectorstore.save_local(str(INDEX_DIR))
    st.cache_resource.clear()
    return True, f"已新增 {len(chunks)} 個向量塊至知識庫"


# ── RAG 鏈 ────────────────────────────────────────

@st.cache_resource(show_spinner=False)
def load_chain(_version: int = 0):
    if not INDEX_DIR.exists():
        return None, None
    embeddings  = OpenAIEmbeddings(model="text-embedding-3-small")
    vectorstore = FAISS.load_local(
        str(INDEX_DIR), embeddings, allow_dangerous_deserialization=True
    )
    retriever = vectorstore.as_retriever(
        search_type="mmr",
        search_kwargs={"k": TOP_K, "fetch_k": 20, "lambda_mult": 0.7},
    )
    llm    = ChatOpenAI(model="gpt-4o-mini", temperature=0.1, max_tokens=800)
    prompt = ChatPromptTemplate.from_messages([
        ("system", SYSTEM_PROMPT),
        MessagesPlaceholder(variable_name="chat_history"),
        ("human", "{input}"),
    ])
    memory = ConversationBufferWindowMemory(
        k=WINDOW_K, memory_key="chat_history",
        return_messages=True, output_key="answer",
    )
    chain  = create_retrieval_chain(retriever, create_stuff_documents_chain(llm, prompt))
    return chain, memory


def format_sources(source_docs: list) -> list:
    seen, lines = set(), []
    for doc in source_docs:
        meta   = doc.metadata
        source = meta.get("source", "未知來源")
        page   = meta.get("page", "")
        key    = f"{source}-{page}"
        if key in seen:
            continue
        seen.add(key)
        label = source
        if page:
            label += f"　第 {page} 頁"
        lines.append(label)
    return lines


# ── 頁面設定 ────────────────────────────────────────
st.set_page_config(
    page_title="教務處法規助理｜國立臺北教育大學",
    page_icon="🎓", layout="centered",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Noto+Sans+TC:wght@400;500;700&display=swap');
html, body, [class*="css"], .stApp { font-family: 'Noto Sans TC', sans-serif !important; }
#MainMenu, footer { visibility: hidden; }
.block-container { padding-top: 1.5rem !important; padding-bottom: 2rem !important; max-width: 780px !important; }
[data-testid="collapsedControl"] { display: flex !important; visibility: visible !important; opacity: 1 !important; }
.ntue-header {
    display: flex; align-items: center; gap: 1rem;
    padding: 1.2rem 1.5rem; background: #0f2d6b;
    border-radius: 12px; margin-bottom: 1.2rem; border-bottom: 3px solid #c8a400;
}
.ntue-badge { width: 46px; height: 46px; background: #c8a400; border-radius: 10px;
    display: flex; align-items: center; justify-content: center; font-size: 22px; flex-shrink: 0; }
.ntue-title h1 { font-size: 1.1rem; font-weight: 700; color: #fff; margin: 0 0 3px; }
.ntue-title p  { font-size: 0.76rem; color: #a8bcd8; margin: 0; }
div[data-testid="stHorizontalBlock"] { gap: 6px !important; }
.tab-btn button {
    background: #f8f9fc !important; border: 1.5px solid #dde3ed !important;
    color: #5a6a85 !important; border-radius: 8px !important;
    font-size: 0.85rem !important; font-weight: 500 !important; padding: 7px 4px !important;
}
.tab-btn-active button {
    background: #0f2d6b !important; border: 1.5px solid #0f2d6b !important;
    color: #fff !important; border-radius: 8px !important;
    font-size: 0.85rem !important; font-weight: 500 !important; padding: 7px 4px !important;
}
.quick-label { font-size: 0.72rem; color: #9ca3af; font-weight: 600;
    letter-spacing: 0.06em; text-transform: uppercase; margin-bottom: 8px; }
div[data-testid="column"] button {
    background: #f0f4ff !important; border: 1.5px solid #d0d9f0 !important;
    color: #2b4a9e !important; border-radius: 8px !important;
    font-size: 0.78rem !important; font-weight: 500 !important;
    padding: 6px 10px !important; white-space: normal !important;
    height: 64px !important; line-height: 1.4 !important;
    display: flex !important; align-items: center !important;
    justify-content: center !important; text-align: center !important;
}
div[data-testid="column"] button:hover { background: #dde6ff !important; border-color: #2b4a9e !important; }
div[data-testid="column"] button p { margin: 0 !important; white-space: normal !important; }
.source-box { background: #f5f7ff; border: 1px solid #d8e0f5; border-left: 3px solid #0f2d6b;
    padding: 8px 12px; border-radius: 0 8px 8px 0; font-size: 0.76rem; color: #4a5568; margin-top: 8px; }
.source-box .source-title { font-weight: 600; margin-bottom: 4px; color: #0f2d6b; }
.source-box .source-line { padding: 2px 0; border-bottom: 1px solid #e8edf8; }
.source-box .source-line:last-child { border-bottom: none; }
[data-testid="stSidebar"] { background: #f8f9fc !important; }
.sb-title { font-size: 0.68rem; font-weight: 700; letter-spacing: 0.08em;
    text-transform: uppercase; color: #9ca3af; margin-bottom: 8px; }
.contact-row { display: flex; justify-content: space-between; align-items: center;
    padding: 5px 0; border-bottom: 1px solid #edf0f7; font-size: 0.82rem; }
.contact-row:last-child { border-bottom: none; }
.contact-ext { color: #0f2d6b; font-weight: 700; }
hr { border-color: #e8ecf4 !important; margin: 0.8rem 0 !important; }
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="ntue-header">
    <div class="ntue-badge">🏫</div>
    <div class="ntue-title">
        <h1>國立臺北教育大學 教務處法規助理</h1>
        <p>查詢學則、法規與辦理流程，每則回答均標註官方來源</p>
    </div>
</div>
""", unsafe_allow_html=True)

# ── Session State ──────────────────────────────────
for key, val in [("messages", []), ("selected_role", "student"), ("pending_input", None), ("index_version", 0)]:
    if key not in st.session_state:
        st.session_state[key] = val

# ── 身份 Tab ──────────────────────────────────────
role_keys = list(QUICK_QUESTIONS.keys())
tab_cols  = st.columns(len(role_keys))
for i, (k, v) in enumerate(QUICK_QUESTIONS.items()):
    is_active = (k == st.session_state.selected_role)
    css_class = "tab-btn-active" if is_active else "tab-btn"
    with tab_cols[i]:
        st.markdown(f'<div class="{css_class}">', unsafe_allow_html=True)
        if st.button(f'{v["icon"]} {v["label"]}', key=f"tab_{k}", use_container_width=True):
            st.session_state.selected_role = k
            st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

role = st.session_state.selected_role

# ── 快速提問 ──────────────────────────────────────
st.markdown('<div class="quick-label">⚡ 快速提問</div>', unsafe_allow_html=True)
qs   = QUICK_QUESTIONS[role]["questions"]
cols = st.columns(len(qs))
for i, q in enumerate(qs):
    if cols[i].button(q, key=f"q_{role}_{i}", use_container_width=True):
        st.session_state.pending_input = q

st.divider()

# ── 載入 RAG 鏈 ───────────────────────────────────
with st.spinner("載入知識庫中..."):
    chain, memory = load_chain(st.session_state.index_version)

if chain is None:
    st.error("尚未建立向量索引，請先執行：python build_index.py")
    st.stop()

# ── 歷史對話 ──────────────────────────────────────
for msg in st.session_state.messages:
    with st.chat_message(msg["role"]):
        st.markdown(msg["content"])

# ── 輸入處理 ──────────────────────────────────────
user_input = st.chat_input("輸入問題，例如：如何申請轉系？") or st.session_state.pending_input
if st.session_state.pending_input:
    st.session_state.pending_input = None

if user_input:
    with st.chat_message("user"):
        st.markdown(user_input)
    st.session_state.messages.append({"role": "user", "content": user_input})

    with st.chat_message("assistant"):
        with st.spinner("查閱法規資料庫中..."):
            try:
                chat_history = memory.load_memory_variables({}).get("chat_history", [])
                result  = chain.invoke({"input": user_input, "chat_history": chat_history})
                answer  = result["answer"]
                sources = format_sources(result.get("context", []))
                memory.save_context({"input": user_input}, {"answer": answer})
                st.markdown(answer)
                if sources:
                    with st.expander("📄 檢視官方參考法規來源"):
                        for src in sources:
                            st.write(f"• {src}")
                st.session_state.messages.append({"role": "assistant", "content": answer})
            except Exception as e:
                st.error("系統發生錯誤，請稍後再試或致電 (02)2732-1104")
                print(f"[ERROR] {e}")

# ── 側邊欄 ────────────────────────────────────────
with st.sidebar:
    st.markdown('<div class="sb-title">📞 聯絡資訊</div>', unsafe_allow_html=True)
    st.markdown("""
<div style="font-size:0.8rem;color:#374151;margin-bottom:8px">
    <strong>總機</strong>：(02) 2732-1104<br>
    <span style="color:#9ca3af;font-size:0.73rem">週一至週五 08:30–17:30</span>
</div>
""", unsafe_allow_html=True)

    contacts = [
        ("教務長室", "82011"), ("招生與宣傳組", "82221"),
        ("註冊組", "82231"), ("課務組", "82016"), ("華語文中心", "82025"),
    ]
    rows = "".join(
        f'<div class="contact-row"><span>{n}</span><span class="contact-ext">分機 {e}</span></div>'
        for n, e in contacts
    )
    st.markdown(f'<div style="margin-bottom:1rem">{rows}</div>', unsafe_allow_html=True)

    st.markdown('<div class="sb-title">🔗 常用系統</div>', unsafe_allow_html=True)
    for label, url in [
        ("iNTUE 校務系統",  "https://nsa.ntue.edu.tw/"),
        ("Moodle 教學平台", "https://md.ntue.edu.tw/"),
        ("校園入口網",       "https://protocol.ntue.edu.tw/"),
        ("計算機中心",       "https://cc.ntue.edu.tw/"),
    ]:
        st.markdown(
            f'<a href="{url}" target="_blank" style="display:block;padding:7px 10px;margin-bottom:4px;'
            f'border-radius:8px;background:#f0f4ff;color:#0f2d6b;font-size:0.82rem;'
            f'font-weight:500;text-decoration:none">🔗 {label}</a>',
            unsafe_allow_html=True,
        )

    st.divider()
    if st.button("🗑️ 清除對話紀錄", use_container_width=True):
        st.session_state.messages = []
        if memory:
            memory.clear()
        st.rerun()

    # ── 管理員上傳區 ──────────────────────────────
    st.divider()
    with st.expander("🔧 管理員"):
        admin_pwd = st.text_input("管理員密碼", type="password", key="admin_pwd")
        if admin_pwd and admin_pwd == os.environ.get("ADMIN_PASSWORD", ""):

            st.success("✅ 已登入")
            st.markdown("**上傳新文件**")

            doc_type = st.radio(
                "文件類型",
                options=["qa", "regulation"],
                format_func=lambda x: "📋 QA 問答集（Word）" if x == "qa" else "📄 法規文件（PDF）",
                horizontal=True,
                key="upload_doc_type",
            )

            # 依類型限制格式
            allowed = ["docx"] if doc_type == "qa" else ["pdf", "txt", "md"]
            type_hint = "DOCX" if doc_type == "qa" else "PDF / TXT / MD"

            uploaded_files = st.file_uploader(
                f"選擇 {type_hint} 檔案（最多 10MB）",
                type=allowed,
                accept_multiple_files=True,
                key="admin_upload",
            )

            if uploaded_files:
                # 驗證並預覽
                valid, invalid = [], []
                for f in uploaded_files:
                    if f.size > 10 * 1024 * 1024:
                        invalid.append(f"❌ {f.name}（超過 10MB）")
                    else:
                        valid.append(f)
                        st.write(f"✅ {f.name}（{f.size/1024:.1f} KB）")
                for msg in invalid:
                    st.warning(msg)

                if valid and st.button("🚀 上傳並更新知識庫", type="primary"):
                    new_docs = []
                    errors   = []

                    with st.spinner("解析文件中..."):
                        REGULATIONS_DIR.mkdir(exist_ok=True)
                        for f in valid:
                            suffix = Path(f.name).suffix.lower()
                            dest   = REGULATIONS_DIR / f.name
                            dest.write_bytes(f.read())

                            try:
                                if suffix == ".pdf":
                                    if is_scanned_pdf(dest):
                                        errors.append(f"⚠️ {f.name} 疑似掃描版 PDF，無法提取文字")
                                        dest.unlink()
                                        continue
                                    new_docs.extend(parse_pdf(dest, doc_type))
                                elif suffix == ".docx":
                                    new_docs.extend(parse_docx(dest, doc_type))
                                elif suffix in (".txt", ".md"):
                                    new_docs.extend(parse_txt(dest, doc_type))
                            except Exception as e:
                                errors.append(f"⚠️ {f.name} 解析失敗：{e}")
                                if dest.exists():
                                    dest.unlink()

                    if errors:
                        for err in errors:
                            st.warning(err)

                    if new_docs:
                        with st.spinner("更新知識庫中（約 10～30 秒）..."):
                            ok, msg = add_to_index(new_docs)
                        if ok:
                            st.session_state.index_version += 1
                            st.session_state.messages = []
                            st.success(f"✅ 知識庫已更新！{msg}")
                            st.rerun()
                        else:
                            st.error(f"❌ 更新失敗：{msg}")
                    elif not errors:
                        st.warning("沒有可處理的文件")

        elif admin_pwd:
            st.error("密碼錯誤")

    st.markdown(
        '<div style="font-size:0.7rem;color:#9ca3af;margin-top:0.5rem;line-height:1.6">'
        '回答依官方法規生成，如有疑義請以原始文件為準。</div>',
        unsafe_allow_html=True,
    )
